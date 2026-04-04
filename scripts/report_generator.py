'''
High-Fidelity Word (.docx) Report Generator for Physics Labs.

Generates documents that strictly follow professional academic formatting:
- Correct fonts (Songti, Heiti, Times New Roman)
- Precise margins and spacing
- Stylized tables with headers and background colors
- Center-aligned OMML formulas with labels
- Integrated charts with captions
'''

import os
import numpy as np
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Import skill components
from data_processor import (
    calculate_mean_expanded,
    calculate_std_dev_expanded,
    bad_value_test,
    calculate_uncertainties_expanded,
    round_physics,
    linear_fit_and_plot
)
from omml_generator import (
    r, frac, rad, ssub, ssup, ssubsup, d, omath, nary, add_formula_para
)

def set_font(run, font_name='宋体', font_size=12, bold=False, italic=False):
    '''Sets font properties for a docx run.'''
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic

def set_cell_background(cell, fill_color="E7E6E6"):
    '''Sets background color for a table cell.'''
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_report(
    experiment_name,
    experiment_date,
    student_name,
    student_class,
    student_id,
    objective,
    principle,
    instruments,
    steps,
    raw_data_tables,
    data_processing_sections,
    error_analysis_discussion,
    conclusion,
    thought_questions,
    output_dir="/home/ubuntu/outputs"
):
    '''Main entry point for generating the high-fidelity report.'''
    doc = Document()

    # --- Page Layout (A4) ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # --- Title Section ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(experiment_name)
    set_font(title_run, '黑体', 16, bold=True)
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = f"姓名: {student_name}  班级: {student_class}  学号: {student_id}"
    info_run = info_p.add_run(info_text)
    set_font(info_run, '宋体', 12)

    # --- I. 原始记录数据 ---
    doc.add_heading("一、原始记录数据", level=1)
    
    for table_info in raw_data_tables:
        if "title" in table_info:
            p = doc.add_paragraph(table_info["title"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.runs[0], '宋体', 10, bold=True)
            
        table = doc.add_table(rows=1, cols=len(table_info["headers"]))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header Row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(table_info["headers"]):
            hdr_cells[i].text = header
            set_cell_background(hdr_cells[i], "D9E1F2") # Light blue
            set_font(hdr_cells[i].paragraphs[0].runs[0], '宋体', 10, bold=True)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Data Rows
        for row_data in table_info["data"]:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                set_font(row_cells[i].paragraphs[0].runs[0], '宋体', 10)
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- II. 数据处理与误差分析 ---
    doc.add_heading("二、数据处理与误差分析", level=1)
    
    for section_title, section_content in data_processing_sections.items():
        doc.add_heading(section_title, level=2)
        for item in section_content:
            if isinstance(item, str):
                p = doc.add_paragraph(item)
                set_font(p.runs[0], '宋体', 12)
            elif isinstance(item, dict) and "omml" in item:
                label = item.get('label')
                add_formula_para(doc, item["omml"], label=label)
            elif isinstance(item, dict) and "image" in item:
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_p.add_run()
                run.add_picture(item["image"], width=Cm(14))
                if "caption" in item:
                    cap_p = doc.add_paragraph(item["caption"])
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_font(cap_p.runs[0], '宋体', 10, italic=True)

    # --- III. 思考题 ---
    doc.add_heading("三、思考题", level=1)
    for q_num, q_content in thought_questions.items():
        p = doc.add_paragraph()
        q_run = p.add_run(f"思考题 {q_num}: {q_content['question']}")
        set_font(q_run, '宋体', 12, bold=True)
        
        ans_p = doc.add_paragraph()
        ans_run = ans_p.add_run(f"解答: {q_content['answer']}")
        set_font(ans_run, '宋体', 12)

    # --- Save ---
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    filename = f"{experiment_name}_实验报告_{student_name}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    print("High-Fidelity Report Generator Loaded.")
    
    # --- Integration Test with Sample Data ---
    exp_name = "薄凸透镜焦距测量实验报告"
    s_name = "邱茂城"
    s_class = "250101"
    s_id = "25009300422"
    
    # 1. Direct Measurement Processing
    data_f = [15.20, 15.18, 15.05, 14.85, 14.95, 15.10, 15.25, 15.05]
    unit = "cm"
    processed_f, mean_f, std_f, bad_log = bad_value_test(data_f, unit)
    delta_x_f, unc_log = calculate_uncertainties_expanded(processed_f, std_f, 0.05, unit)
    res_str, _, round_log = round_physics(mean_f, delta_x_f, unit)
    
    # Formula Example: Mean
    # f_bar = (1/n) * Sum(f_i)
    mean_formula = omath([
        ssup(r("f", italic=True), r("\u0305")), r(" = "),
        frac(r("1"), r(str(len(processed_f)))),
        nary("\u03a3", lower=ssub(r("i"), r("1")), upper=r(str(len(processed_f))), content=ssub(r("f", italic=True), r("i")))
    ])
    
    # Formula Example: Result
    # f = (mean +/- delta) cm
    result_formula = omath([
        r("f", italic=True), r(" = "), r(res_str)
    ])

    data_sections = {
        "1. 自准法（直接测量）": [
            "(1) 计算算术均值:",
            {"omml": mean_formula, "label": "1-1"},
            f"均值计算结果: {mean_f:.3f} {unit}",
            "(2) 坏值检验 (3σ准则):",
            bad_log,
            "(3) 不确定度计算与修约:",
            unc_log,
            round_log,
            {"omml": result_formula, "label": "1-2"}
        ]
    }
    
    raw_tables = [
        {
            "title": "表1 自准法测量凸透镜焦距原始数据",
            "headers": ["次数 n", "1", "2", "3", "4", "5", "6", "7", "8"],
            "data": [["f / cm"] + [str(x) for x in data_f]],
            "col_widths": [2] + [1.5]*8
        }
    ]
    
    questions = {
        "1": {"question": "如何减小测量误差？", "answer": "多次测量取平均值，严格遵守操作规程。"}
    }

    try:
        path = create_report(exp_name, "2026-04-04", s_name, s_class, s_id, 
                            "测量透镜焦距", "透镜成像公式", {}, [], 
                            raw_tables, data_sections, "误差主要来源于...", "结论良好", questions)
        print(f"Test report generated at: {path}")
    except Exception as e:
        print(f"Error: {e}")
