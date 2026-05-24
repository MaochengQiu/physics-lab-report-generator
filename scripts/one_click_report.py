"""
Weak-model friendly one-click report generator.

Usage:
1. Edit REPORT_CONFIG only.
2. Run: python scripts/one_click_report.py
3. Read the final DOCX_STATUS / DOCX_PATH line.

This file intentionally hides the lower-level data processing, plotting, OMML,
and docx assembly details so less capable models only need to fill structured data.
"""

from __future__ import annotations

import math
import os
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from scipy import stats

try:
    from omml_generator import add_formula_para, omath, r
except ImportError:
    from scripts.omml_generator import add_formula_para, omath, r


REPORT_CONFIG: dict[str, Any] = {
    "experiment_name": "示例：弹簧振子周期测量实验报告",
    "student": {
        "name": "张三",
        "class": "物理2501",
        "id": "2025000000",
        "date": "2026-05-24",
    },
    "objective": [
        "学习等精度多次测量的数据处理方法。",
        "掌握不确定度合成、物理修约与实验图像分析。",
    ],
    "principle": [
        "对同一物理量进行多次等精度测量，先进行坏值检验，再计算均值、样本标准差和不确定度。",
        "若实验变量存在近似线性关系，使用最小二乘线性拟合判断关系并提取斜率、截距和相关系数。",
    ],
    "instruments": [
        {"method": "周期测量法", "name": "秒表", "delta": 0.01, "unit": "s"},
        {"method": "长度测量法", "name": "钢尺", "delta": 0.05, "unit": "cm"},
    ],
    "steps": [
        "检查实验仪器并记录仪器精度。",
        "按实验要求完成多次测量并记录原始数据。",
        "进行数据处理、作图和误差分析。",
    ],
    "measurements": [
        {
            "method": "周期测量法",
            "symbol": "T",
            "name": "周期",
            "unit": "s",
            "delta_instrument": 0.01,
            "data": [1.42, 1.43, 1.41, 1.44, 1.42, 1.43],
        },
        {
            "method": "长度测量法",
            "symbol": "L",
            "name": "长度",
            "unit": "cm",
            "delta_instrument": 0.05,
            "data": [20.12, 20.10, 20.14, 20.11, 20.13, 20.12],
        },
    ],
    "plot": {
        "mode": "linear",  # linear or trend
        "title": "T-L 关系图",
        "x_label": "L / cm",
        "y_label": "T / s",
        "x_data": [20.10, 21.20, 22.30, 23.40, 24.50, 25.60],
        "y_data": [1.41, 1.45, 1.49, 1.53, 1.57, 1.61],
        "gamma0": 0.75,
    },
    "error_analysis": [
        "随机误差主要来自读数波动和操作重复性差异。",
        "系统误差主要来自仪器零点、刻度精度和实验装置安装偏差。",
    ],
    "conclusion": [
        "各直接测量量均完成坏值检验、不确定度计算和物理修约。",
        "实验图像已插入报告，用于辅助判断变量关系与结果可靠性。",
    ],
    "thought_questions": [
        {"question": "如何减小随机误差？", "answer": "增加测量次数，保持操作条件一致，并对数据进行坏值检验。"},
    ],
    "output_dir": "outputs",
}


class ReportError(RuntimeError):
    def __init__(self, reason: str, required_inputs: list[str] | None = None):
        super().__init__(reason)
        self.reason = reason
        self.required_inputs = required_inputs or []


def set_font(run, font_name="宋体", font_size=12, bold=False, italic=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic


def add_paragraph(doc: Document, text: str, font_size=12, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, font_size=font_size, bold=bold, italic=italic)
    p.paragraph_format.first_line_indent = Cm(0.75) if align is None else None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    return p


def set_cell_background(cell, fill_color="E7E6E6"):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[Any]]):
    add_paragraph(doc, title, font_size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_background(cell, "D9E1F2")
        set_font(cell.paragraphs[0].runs[0], font_size=10, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_font(cells[i].paragraphs[0].runs[0], font_size=10)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def validate_config(config: dict[str, Any]) -> None:
    required_top = ["experiment_name", "student", "measurements", "plot", "output_dir"]
    missing = [key for key in required_top if not config.get(key)]
    if missing:
        raise ReportError("配置缺少顶层字段。", missing)

    for index, measurement in enumerate(config["measurements"], start=1):
        required = ["method", "symbol", "unit", "delta_instrument", "data"]
        missing = [key for key in required if key not in measurement or measurement[key] in (None, "", [])]
        if missing:
            raise ReportError(f"第 {index} 组测量缺少必要字段。", [f"measurements[{index}].{key}" for key in missing])
        if len(measurement["data"]) < 2:
            raise ReportError(f"第 {index} 组测量数据少于 2 个，无法计算样本标准差。", [f"measurements[{index}].data"])

    plot = config["plot"]
    for key in ["x_data", "y_data", "x_label", "y_label", "title"]:
        if key not in plot or plot[key] in (None, "", []):
            raise ReportError("作图门失败：缺少作图字段。", [f"plot.{key}"])
    if len(plot["x_data"]) != len(plot["y_data"]) or len(plot["x_data"]) < 2:
        raise ReportError("作图门失败：x/y 数据长度不一致或点数不足。", ["plot.x_data", "plot.y_data"])


def mean_expansion(data: list[float], symbol: str, unit: str) -> tuple[float, str]:
    n = len(data)
    total = sum(data)
    mean = total / n
    terms = "+".join(format_number(x) for x in data)
    return mean, f"{symbol}̄ = 1/{n} × ({terms}) = {total:.5g}/{n} = {mean:.5g} {unit}"


def sample_std(data: list[float], mean: float, symbol: str, unit: str) -> tuple[float, float, str]:
    n = len(data)
    residual_sum = sum((x - mean) ** 2 for x in data)
    std = math.sqrt(residual_sum / (n - 1))
    return std, residual_sum, f"σ_{symbol} = sqrt({residual_sum:.6g}/{n - 1}) = {std:.6g} {unit}"


def three_sigma_loop(data: list[float], unit: str) -> tuple[list[float], float, float, list[str]]:
    current = [float(x) for x in data]
    logs: list[str] = []
    removed = 0

    while len(current) >= 3:
        mean = sum(current) / len(current)
        std = math.sqrt(sum((x - mean) ** 2 for x in current) / (len(current) - 1))
        residuals = [abs(x - mean) for x in current]
        logs.append(f"本轮均值 = {mean:.6g} {unit}，σ = {std:.6g} {unit}。")
        logs.append("残差 |xi-x̄| = " + ", ".join(f"{v:.6g}" for v in residuals) + f" {unit}")

        if std == 0:
            logs.append("所有数据残差为 0，无坏值。")
            break

        max_residual = max(residuals)
        if max_residual > 3 * std:
            index = residuals.index(max_residual)
            value = current.pop(index)
            removed += 1
            logs.append(f"第 {index + 1} 个数据 {value:g} {unit} 满足 |xi-x̄| > 3σ，已剔除并重新计算。")
        else:
            logs.append("所有数据均满足 |xi-x̄| ≤ 3σ，无坏值。")
            break

    mean = sum(current) / len(current)
    std = math.sqrt(sum((x - mean) ** 2 for x in current) / (len(current) - 1)) if len(current) > 1 else 0.0
    logs.append(f"坏值检验结束：共剔除 {removed} 个坏值。")
    return current, mean, std, logs


def calculate_uncertainty(data: list[float], std: float, delta_instrument: float, unit: str) -> tuple[float, float, float, list[str]]:
    n = len(data)
    delta_a = std / math.sqrt(n)
    delta_b = delta_instrument / math.sqrt(3)
    delta = math.sqrt(delta_a**2 + delta_b**2)
    logs = [
        f"ΔA = σ/sqrt(n) = {std:.6g}/sqrt({n}) = {delta_a:.6g} {unit}",
        f"ΔB = Δ仪/sqrt(3) = {delta_instrument:.6g}/sqrt(3) = {delta_b:.6g} {unit}",
        f"Δ = sqrt(ΔA²+ΔB²) = sqrt({delta_a**2:.6g}+{delta_b**2:.6g}) = {delta:.6g} {unit}",
    ]
    return delta_a, delta_b, delta, logs


def ceil_to_sig(value: float, sig_digits: int) -> float:
    if value <= 0 or not math.isfinite(value):
        return value
    exponent = math.floor(math.log10(abs(value)))
    factor = 10 ** (sig_digits - 1 - exponent)
    return math.ceil(value * factor - 1e-12) / factor


def decimals_for(value: float) -> int:
    if value == 0:
        return 0
    return max(0, -math.floor(math.log10(abs(value))))


def round_result(mean: float, uncertainty: float, unit: str) -> tuple[str, float, float, str]:
    if uncertainty <= 0 or not math.isfinite(uncertainty):
        raise ReportError("不确定度无效，无法修约。", ["measurement.delta_instrument", "measurement.data"])
    first_digit = int(abs(uncertainty) / (10 ** math.floor(math.log10(abs(uncertainty)))))
    sig_digits = 2 if first_digit in (1, 2) else 1
    rounded_uncertainty = ceil_to_sig(uncertainty, sig_digits)
    decimals = decimals_for(rounded_uncertainty)
    rounded_mean = round(mean, decimals)
    result = f"({rounded_mean:.{decimals}f} ± {rounded_uncertainty:.{decimals}f}) {unit}"
    log = f"ΔX 首位为 {first_digit}，保留 {sig_digits} 位有效数字并只进不舍；x̄ 末位与 ΔX 对齐，结果为 {result}。"
    return result, rounded_mean, rounded_uncertainty, log


def format_number(value: float) -> str:
    return f"{value:.6g}"


def process_measurement(measurement: dict[str, Any]) -> dict[str, Any]:
    data = [float(x) for x in measurement["data"]]
    unit = measurement["unit"]
    symbol = measurement["symbol"]
    processed, mean_after_test, std_after_test, bad_logs = three_sigma_loop(data, unit)
    mean, mean_log = mean_expansion(processed, symbol, unit)
    std, residual_sum, std_log = sample_std(processed, mean, symbol, unit)
    delta_a, delta_b, delta, uncertainty_logs = calculate_uncertainty(processed, std, float(measurement["delta_instrument"]), unit)
    final_result, rounded_mean, rounded_delta, rounding_log = round_result(mean, delta, unit)

    return {
        "method": measurement["method"],
        "symbol": symbol,
        "name": measurement.get("name", symbol),
        "unit": unit,
        "raw_data": data,
        "processed_data": processed,
        "mean": mean,
        "std": std,
        "residual_sum": residual_sum,
        "delta_a": delta_a,
        "delta_b": delta_b,
        "delta": delta,
        "rounded_mean": rounded_mean,
        "rounded_delta": rounded_delta,
        "final_result": final_result,
        "mean_log": mean_log,
        "std_log": std_log,
        "bad_logs": bad_logs,
        "uncertainty_logs": uncertainty_logs,
        "rounding_log": rounding_log,
    }


def make_plot(plot_config: dict[str, Any], output_dir: Path) -> dict[str, Any]:
    x = np.array(plot_config["x_data"], dtype=float)
    y = np.array(plot_config["y_data"], dtype=float)
    mode = plot_config.get("mode", "linear")
    title = plot_config.get("title", "实验数据图")
    path = output_dir / "experiment_plot.png"

    plt.figure(figsize=(8, 5), dpi=150)
    plt.scatter(x, y, color="#FF4C4C", s=64, label="实验数据")

    result: dict[str, Any] = {"mode": mode, "path": str(path)}
    if mode == "linear":
        slope, intercept, r_value, _p_value, std_err = stats.linregress(x, y)
        x_fit = np.linspace(float(np.min(x)), float(np.max(x)), 200)
        y_fit = slope * x_fit + intercept
        plt.plot(x_fit, y_fit, color="#0066FF", linewidth=2.5, label="线性拟合")
        gamma0 = float(plot_config.get("gamma0", 0.75))
        relation = "满足线性关系" if abs(r_value) > gamma0 else "线性关系较弱"
        text = f"k = {slope:.5g}\nb = {intercept:.3f}\nr = {r_value:.5f}\n{relation}"
        plt.text(
            0.97,
            0.97,
            text,
            transform=plt.gca().transAxes,
            ha="right",
            va="top",
            bbox={"boxstyle": "round", "facecolor": "white", "edgecolor": "#FFD966"},
        )
        result.update({
            "slope": slope,
            "intercept": intercept,
            "r": r_value,
            "std_err": std_err,
            "equation": f"y = {slope:.5g}x + {intercept:.3f}",
            "relation": relation,
        })
    else:
        order = np.argsort(x)
        plt.plot(x[order], y[order], color="#0066FF", linewidth=2.0, label="趋势线")
        result.update({"relation": "已生成趋势图"})

    plt.title(title, fontsize=14)
    plt.xlabel(plot_config["x_label"])
    plt.ylabel(plot_config["y_label"])
    plt.grid(True, linestyle="--", color="#BBBBBB", alpha=0.7)
    plt.legend(loc="upper left")
    plt.tight_layout()
    plt.savefig(path, dpi=300)
    plt.close()

    if not path.exists():
        raise ReportError("作图门失败：图片文件未成功保存。", ["plot"])
    return result


def add_simple_formula(doc: Document, formula_text: str, label: str):
    add_formula_para(doc, omath(r(formula_text)), label=label)


def build_docx(config: dict[str, Any], processed: list[dict[str, Any]], plot_result: dict[str, Any], output_dir: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(config["experiment_name"])
    set_font(title_run, "黑体", 16, bold=True)

    student = config["student"]
    info = f"姓名：{student.get('name', '')}    班级：{student.get('class', '')}    学号：{student.get('id', '')}    日期：{student.get('date', '')}"
    add_paragraph(doc, info, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_heading("一、实验目的", level=1)
    for item in config.get("objective", []):
        add_paragraph(doc, item)

    doc.add_heading("二、实验原理", level=1)
    for item in config.get("principle", []):
        add_paragraph(doc, item)
    add_simple_formula(doc, "x̄ = (1/n)Σxi", "1")
    add_simple_formula(doc, "σ = sqrt(Σ(xi-x̄)²/(n-1))", "2")
    add_simple_formula(doc, "ΔX = sqrt(ΔA²+ΔB²)", "3")

    doc.add_heading("三、实验仪器", level=1)
    instrument_rows = [[x.get("method", ""), x.get("name", ""), x.get("delta", ""), x.get("unit", "")] for x in config.get("instruments", [])]
    add_table(doc, "表1 实验仪器与仪器误差", ["方法", "仪器", "Δ仪", "单位"], instrument_rows)

    doc.add_heading("四、实验步骤", level=1)
    for index, item in enumerate(config.get("steps", []), start=1):
        add_paragraph(doc, f"{index}. {item}")

    doc.add_heading("五、数据记录与处理", level=1)
    for index, result in enumerate(processed, start=1):
        doc.add_heading(f"5.{index} {result['method']}", level=2)
        headers = ["序号"] + [str(i) for i in range(1, len(result["raw_data"]) + 1)]
        rows = [[f"{result['symbol']} / {result['unit']}"] + [format_number(x) for x in result["raw_data"]]]
        add_table(doc, f"表{index + 1} {result['method']}原始数据", headers, rows)

        doc.add_heading("（1）算术均值", level=3)
        add_simple_formula(doc, f"{result['symbol']}̄ = (1/n)Σ{result['symbol']}i", f"5-{index}-1")
        add_paragraph(doc, "代入展开式：" + result["mean_log"])

        doc.add_heading("（2）样本标准差", level=3)
        add_simple_formula(doc, f"σ = sqrt(Σ({result['symbol']}i-{result['symbol']}̄)²/(n-1))", f"5-{index}-2")
        add_paragraph(doc, "残差平方和代入：" + result["std_log"])

        doc.add_heading("（3）坏值检验（3σ准则）", level=3)
        add_simple_formula(doc, "|xi-x̄| ≤ 3σ", f"5-{index}-3")
        for log in result["bad_logs"]:
            add_paragraph(doc, log)

        doc.add_heading("（4）不确定度计算", level=3)
        add_simple_formula(doc, "ΔA = σ/sqrt(n)", f"5-{index}-4")
        add_simple_formula(doc, "ΔB = Δ仪/sqrt(3)", f"5-{index}-5")
        add_simple_formula(doc, "ΔX = sqrt(ΔA²+ΔB²)", f"5-{index}-6")
        for log in result["uncertainty_logs"]:
            add_paragraph(doc, "代入计算：" + log)

        doc.add_heading("（5）修约与结果表示", level=3)
        add_paragraph(doc, result["rounding_log"])
        add_paragraph(doc, f"最终结果：{result['symbol']} = {result['final_result']}", bold=True)

    doc.add_heading("六、实验图像与拟合结果", level=1)
    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.add_run().add_picture(plot_result["path"], width=Cm(14))
    add_paragraph(doc, f"图1 {config['plot'].get('title', '实验数据图')}", font_size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    if plot_result.get("mode") == "linear":
        rows = [[
            f"{plot_result['slope']:.5g}",
            f"{plot_result['intercept']:.3f}",
            f"{plot_result['r']:.5f}",
            plot_result["equation"],
            plot_result["relation"],
        ]]
        add_table(doc, "表 拟合结果汇总", ["斜率 k", "截距 b", "相关系数 r", "拟合方程", "结论"], rows)
    else:
        add_paragraph(doc, plot_result.get("relation", "已生成趋势图"))

    doc.add_heading("七、误差分析与讨论", level=1)
    for item in config.get("error_analysis", []):
        add_paragraph(doc, item)

    doc.add_heading("八、实验结论", level=1)
    for item in config.get("conclusion", []):
        add_paragraph(doc, item)
    for result in processed:
        add_paragraph(doc, f"{result['symbol']} = {result['final_result']}", bold=True)

    doc.add_heading("九、思考题", level=1)
    for index, item in enumerate(config.get("thought_questions", []), start=1):
        add_paragraph(doc, f"思考题 {index}：{item.get('question', '')}", bold=True)
        add_paragraph(doc, f"解答：{item.get('answer', '')}")

    safe_name = "".join(ch for ch in config["experiment_name"] if ch not in r'<>:"/\\|?*').strip()
    student_name = student.get("name", "学生")
    path = output_dir / f"{safe_name}_{student_name}.docx"
    doc.save(path)
    if not path.exists():
        raise ReportError("保存门失败：docx 文件未落盘。", ["output_dir"])
    return path


def run(config: dict[str, Any] = REPORT_CONFIG) -> Path:
    validate_config(config)
    output_dir = Path(config.get("output_dir", "outputs")).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    processed = [process_measurement(item) for item in config["measurements"]]
    plot_result = make_plot(config["plot"], output_dir)
    docx_path = build_docx(config, processed, plot_result, output_dir)
    return docx_path


if __name__ == "__main__":
    try:
        output_path = run(REPORT_CONFIG)
        print("DOCX_STATUS: SUCCESS")
        print(f"DOCX_PATH: {output_path}")
        print("SUMMARY: 已完成数据处理、强制作图、OMML公式段落和Word文档生成。")
    except ReportError as exc:
        print("DOCX_STATUS: FAILED")
        print(f"FAILED_REASON: {exc.reason}")
        print(f"REQUIRED_INPUTS: {exc.required_inputs}")
    except Exception as exc:
        print("DOCX_STATUS: FAILED")
        print(f"FAILED_REASON: 未预期错误：{exc}")
        print("REQUIRED_INPUTS: ['检查Python依赖：python-docx, matplotlib, numpy, scipy']")
